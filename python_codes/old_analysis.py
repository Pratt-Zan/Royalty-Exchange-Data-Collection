import os
from bs4 import BeautifulSoup
from docx import Document

def html_to_word(html_file_path, output_docx_path):
    # 1. 读取 HTML 文件
    try:
        with open(html_file_path, 'r', encoding='utf-8') as file:
            html_content = file.read()
    except Exception as e:
        print(f"读取文件失败: {e}")
        return

    # 2. 使用 BeautifulSoup 解析 HTML
    soup = BeautifulSoup(html_content, 'html.parser')

    # 3. 清洗数据：移除所有的 script 和 style 标签
    for script_or_style in soup(['script', 'style']):
        script_or_style.decompose()

    # 4. 提取纯文本
    # separator='\n' 保证不同标签的文字之间有换行
    # strip=True 会去除首尾多余的空白字符
    text_content = soup.get_text(separator='\n', strip=True)

    # 5. 进一步清理空行（可选，让 Word 排版更好看）
    lines = (line.strip() for line in text_content.splitlines())
    clean_text = '\n'.join(line for line in lines if line)

    # 6. 写入 Word 文档
    doc = Document()
    doc.add_heading('提取的网页文本', level=1)
    
    # 按段落写入
    for paragraph in clean_text.split('\n'):
        doc.add_paragraph(paragraph)

    # 7. 保存文件
    try:
        doc.save(output_docx_path)
        print(f"提取成功！文件已保存至: {output_docx_path}")
    except Exception as e:
        print(f"保存 Word 文件失败: {e}")

# 测试运行
if __name__ == "__main__":
    # 替换为你实际的 HTML 文件路径和想要的输出路径
    input_file = "C:\\Users\\Pratt\\Desktop\\HKUST-RA\\Data Collection Royalty exchange\\resources\\old\\1-salsa-hit-by-gilberto-santa-rosa.html"
    output_file = "C:\\Users\\Pratt\\Desktop\\HKUST-RA\\Data Collection Royalty exchange\\analysis\\old\\1-salsa-hit-by-gilberto-santa-rosa.docx"
    
    # 确保输入文件存在再执行
    if os.path.exists(input_file):
        html_to_word(input_file, output_file)
    else:
        print(f"找不到文件: {input_file}")